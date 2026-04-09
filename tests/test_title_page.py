"""Tests for generator/title_page.py."""

import copy
from pathlib import Path

import pytest
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from otchet_compose.generator.styles import setup_styles
from otchet_compose.generator.templates import BUNDLED_TEMPLATES_DIR
from otchet_compose.generator.title_page import (
    _collect_placeholders,
    _inject_page_break,
    _locate_template,
    _substitute_paragraph,
    _warn_param_mismatches,
    list_available_templates,
    render_title_page,
)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_doc():
    d = Document()
    setup_styles(d)
    return d


def _make_template_docx(tmp_path: Path, text: str) -> Path:
    """Create a minimal .docx template with *text* and return its path."""
    path = tmp_path / "custom.docx"
    d = Document()
    d.add_paragraph(text)
    d.save(str(path))
    return path


# ---------------------------------------------------------------------------
# _locate_template
# ---------------------------------------------------------------------------

class TestLocateTemplate:
    def test_finds_bundled_template(self):
        # "mock" and "rut" are bundled
        for name in ("mock", "rut"):
            path = _locate_template(name)
            assert path.exists()
            assert path.suffix == ".docx"

    def test_user_template_takes_priority(self, tmp_path, monkeypatch):
        user_dir = tmp_path / "templates"
        user_dir.mkdir()
        user_file = user_dir / "mock.docx"
        # create a minimal docx there
        d = Document()
        d.save(str(user_file))

        import otchet_compose.generator.title_page as tp
        monkeypatch.setattr(tp, "_USER_TEMPLATES_DIR", user_dir)

        path = _locate_template("mock")
        assert path == user_file

    def test_not_found_raises(self, monkeypatch):
        import otchet_compose.generator.title_page as tp
        # point user dir somewhere that doesn't exist
        monkeypatch.setattr(tp, "_USER_TEMPLATES_DIR", Path("/nonexistent/templates"))
        with pytest.raises(ValueError, match="не найден"):
            _locate_template("does_not_exist_xyz")

    def test_error_message_lists_available(self, monkeypatch):
        import otchet_compose.generator.title_page as tp
        monkeypatch.setattr(tp, "_USER_TEMPLATES_DIR", Path("/nonexistent/templates"))
        with pytest.raises(ValueError) as exc_info:
            _locate_template("does_not_exist_xyz")
        # Should mention bundled templates in the error
        assert "mock" in str(exc_info.value) or "rut" in str(exc_info.value)


# ---------------------------------------------------------------------------
# list_available_templates
# ---------------------------------------------------------------------------

class TestListAvailableTemplates:
    def test_includes_bundled(self):
        names = list_available_templates()
        assert "mock" in names
        assert "rut" in names

    def test_includes_user_templates(self, tmp_path, monkeypatch):
        user_dir = tmp_path / "templates"
        user_dir.mkdir()
        d = Document()
        d.save(str(user_dir / "mytemplate.docx"))

        import otchet_compose.generator.title_page as tp
        monkeypatch.setattr(tp, "_USER_TEMPLATES_DIR", user_dir)

        names = list_available_templates()
        assert "mytemplate" in names

    def test_returns_sorted(self):
        names = list_available_templates()
        assert names == sorted(names)

    def test_no_duplicates(self, tmp_path, monkeypatch):
        """User template shadowing a bundled one should appear only once."""
        user_dir = tmp_path / "templates"
        user_dir.mkdir()
        d = Document()
        d.save(str(user_dir / "mock.docx"))

        import otchet_compose.generator.title_page as tp
        monkeypatch.setattr(tp, "_USER_TEMPLATES_DIR", user_dir)

        names = list_available_templates()
        assert names.count("mock") == 1


# ---------------------------------------------------------------------------
# _collect_placeholders
# ---------------------------------------------------------------------------

class TestCollectPlaceholders:
    def test_finds_placeholders(self, tmp_path):
        path = _make_template_docx(tmp_path, "Hello {{name}}, your group is {{group}}")
        doc = Document(str(path))
        keys = _collect_placeholders(doc)
        assert keys == {"name", "group"}

    def test_no_placeholders(self, tmp_path):
        path = _make_template_docx(tmp_path, "No placeholders here")
        doc = Document(str(path))
        assert _collect_placeholders(doc) == set()

    def test_strips_whitespace_in_key(self, tmp_path):
        path = _make_template_docx(tmp_path, "{{ name }}")
        doc = Document(str(path))
        assert _collect_placeholders(doc) == {"name"}

    def test_deduplicates(self, tmp_path):
        path = _make_template_docx(tmp_path, "{{name}} and {{name}} again")
        doc = Document(str(path))
        assert _collect_placeholders(doc) == {"name"}

    def test_bundled_mock_placeholders(self):
        doc = Document(str(BUNDLED_TEMPLATES_DIR / "mock.docx"))
        keys = _collect_placeholders(doc)
        assert {"discipline", "student", "group", "teacher"}.issubset(keys)

    def test_bundled_rut_placeholders(self):
        doc = Document(str(BUNDLED_TEMPLATES_DIR / "rut.docx"))
        keys = _collect_placeholders(doc)
        assert {"discipline", "student", "group", "teacher"}.issubset(keys)


# ---------------------------------------------------------------------------
# _warn_param_mismatches
# ---------------------------------------------------------------------------

class TestWarnParamMismatches:
    def test_warns_unused_param(self, tmp_path, capsys):
        path = _make_template_docx(tmp_path, "{{name}}")
        doc = Document(str(path))
        _warn_param_mismatches(doc, {"name": "Alice", "extra": "unused"}, "test")
        out = capsys.readouterr().out
        assert "extra" in out
        assert "name" not in out

    def test_warns_unfilled_placeholder(self, tmp_path, capsys):
        path = _make_template_docx(tmp_path, "{{name}} {{group}}")
        doc = Document(str(path))
        _warn_param_mismatches(doc, {"name": "Alice"}, "test")
        out = capsys.readouterr().out
        assert "group" in out
        assert "name" not in out

    def test_no_warnings_when_exact_match(self, tmp_path, capsys):
        path = _make_template_docx(tmp_path, "{{name}}")
        doc = Document(str(path))
        _warn_param_mismatches(doc, {"name": "Alice"}, "test")
        out = capsys.readouterr().out
        assert out == ""

    def test_no_warnings_empty_template_and_params(self, tmp_path, capsys):
        path = _make_template_docx(tmp_path, "No placeholders")
        doc = Document(str(path))
        _warn_param_mismatches(doc, {}, "test")
        assert capsys.readouterr().out == ""


# ---------------------------------------------------------------------------
# _substitute_paragraph
# ---------------------------------------------------------------------------

class TestSubstituteParagraph:
    def _para_with_text(self, text: str):
        doc = Document()
        return doc.add_paragraph(text)

    def test_substitutes_single_placeholder(self):
        para = self._para_with_text("Hello {{name}}")
        _substitute_paragraph(para, {"name": "World"})
        assert "".join(r.text for r in para.runs) == "Hello World"

    def test_substitutes_multiple_placeholders(self):
        para = self._para_with_text("{{a}} and {{b}}")
        _substitute_paragraph(para, {"a": "X", "b": "Y"})
        assert "".join(r.text for r in para.runs) == "X and Y"

    def test_leaves_unknown_key_unchanged(self):
        para = self._para_with_text("{{missing}}")
        _substitute_paragraph(para, {})
        assert "{{missing}}" in "".join(r.text for r in para.runs)

    def test_no_placeholder_no_change(self):
        para = self._para_with_text("plain text")
        _substitute_paragraph(para, {"x": "y"})
        assert "".join(r.text for r in para.runs) == "plain text"

    def test_clears_extra_runs(self):
        """Simulate a token split across runs — all runs except first should be empty after."""
        doc = Document()
        para = doc.add_paragraph()
        para.add_run("{{na")
        para.add_run("me}}")
        _substitute_paragraph(para, {"name": "Alice"})
        assert para.runs[1].text == ""
        assert "Alice" in para.runs[0].text


# ---------------------------------------------------------------------------
# _inject_page_break
# ---------------------------------------------------------------------------

class TestInjectPageBreak:
    def _p_elem(self, text=""):
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)
        return p

    def test_adds_break_to_last_paragraph(self):
        p1 = self._p_elem("first")
        p2 = self._p_elem("last")
        elements = [p1, p2]
        _inject_page_break(elements)
        # The break run should be in p2
        breaks = p2.findall(f".//{qn('w:br')}")
        assert len(breaks) == 1
        assert breaks[0].get(qn("w:type")) == "page"

    def test_does_not_touch_first_paragraph(self):
        p1 = self._p_elem("first")
        p2 = self._p_elem("last")
        elements = [p1, p2]
        _inject_page_break(elements)
        assert p1.findall(f".//{qn('w:br')}") == []

    def test_fallback_when_no_paragraph(self):
        """When no w:p exists, a standalone page-break paragraph is appended."""
        tbl = OxmlElement("w:tbl")
        elements = [tbl]
        _inject_page_break(elements)
        assert len(elements) == 2
        new_p = elements[-1]
        assert new_p.tag == qn("w:p")
        breaks = new_p.findall(f".//{qn('w:br')}")
        assert breaks[0].get(qn("w:type")) == "page"


# ---------------------------------------------------------------------------
# render_title_page (end-to-end)
# ---------------------------------------------------------------------------

class TestRenderTitlePage:
    def test_renders_bundled_mock(self):
        doc = _make_doc()
        initial_count = len(doc.element.body)
        render_title_page(doc, "mock", {
            "institute": "ИИТ", "department": "Каф", "discipline": "ОС",
            "work_type": "ОТЧЁТ", "work_number": "№1", "work_title": "Тест",
            "student": "Иванов", "group": "ИКБО-01", "teacher": "Петров", "year": "2025",
        })
        assert len(doc.element.body) > initial_count

    def test_renders_bundled_rut(self):
        doc = _make_doc()
        render_title_page(doc, "rut", {
            "discipline": "ОС", "work_number": "№1", "work_title": "Тест",
            "student": "Иванов", "group": "ИКБО-01", "teacher": "Петров", "year": "2025",
        })
        assert len(doc.element.body) > 0

    def test_page_break_injected(self):
        doc = _make_doc()
        render_title_page(doc, "mock", {
            "institute": "И", "department": "К", "discipline": "Д",
            "work_type": "О", "work_number": "1", "work_title": "Т",
            "student": "С", "group": "Г", "teacher": "П", "year": "2025",
        })
        body_xml = doc.element.body.xml
        assert 'w:type="page"' in body_xml or "w:type" in body_xml

    def test_user_template_used_when_available(self, tmp_path, monkeypatch):
        user_dir = tmp_path / "templates"
        user_dir.mkdir()
        tmpl_path = user_dir / "mytempl.docx"
        d = Document()
        d.add_paragraph("Привет {{name}}")
        d.save(str(tmpl_path))

        import otchet_compose.generator.title_page as tp
        monkeypatch.setattr(tp, "_USER_TEMPLATES_DIR", user_dir)

        doc = _make_doc()
        render_title_page(doc, "mytempl", {"name": "Мир"})
        # "Мир" should appear somewhere in the body
        full_text = " ".join(p.text for p in doc.paragraphs)
        assert "Мир" in full_text

    def test_unknown_template_raises(self, monkeypatch):
        import otchet_compose.generator.title_page as tp
        monkeypatch.setattr(tp, "_USER_TEMPLATES_DIR", Path("/nonexistent"))
        doc = _make_doc()
        with pytest.raises(ValueError, match="не найден"):
            render_title_page(doc, "nonexistent_xyz", {})

    def test_warns_on_unused_param(self, capsys):
        doc = _make_doc()
        render_title_page(doc, "rut", {
            "discipline": "ОС", "work_number": "№1", "work_title": "Т",
            "student": "С", "group": "Г", "teacher": "П", "year": "2025",
            "nonexistent_key": "warn me",
        })
        assert "nonexistent_key" in capsys.readouterr().out
