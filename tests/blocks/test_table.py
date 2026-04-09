"""Tests for the table block handler."""

from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from otchet_compose.generator.blocks._base import RenderContext
from otchet_compose.generator.blocks.table import TableHandler
from otchet_compose.generator.styles import setup_styles

BASE_DIR = Path(".")


@pytest.fixture()
def handler():
    return TableHandler()


@pytest.fixture()
def doc():
    d = Document()
    setup_styles(d)
    return d


class TestTableValidate:
    def test_valid_block(self, handler):
        block = {
            "type": "table",
            "caption": "Some table",
            "headers": ["A", "B"],
            "rows": [["1", "2"], ["3", "4"]],
        }
        result = handler.validate(block, 1, BASE_DIR)
        assert result["caption"] == "Some table"
        assert result["headers"] == ["A", "B"]
        assert result["rows"] == [["1", "2"], ["3", "4"]]

    def test_caption_is_stripped(self, handler):
        block = {"type": "table", "caption": "  cap  ", "headers": ["A"], "rows": [["x"]]}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["caption"] == "cap"

    def test_headers_are_stripped(self, handler):
        block = {"type": "table", "caption": "c", "headers": ["  H  "], "rows": [["x"]]}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["headers"] == ["H"]

    def test_missing_caption_raises(self, handler):
        block = {"type": "table", "headers": ["A"], "rows": [["x"]]}
        with pytest.raises(ValueError, match="caption"):
            handler.validate(block, 1, BASE_DIR)

    def test_empty_caption_raises(self, handler):
        block = {"type": "table", "caption": "  ", "headers": ["A"], "rows": [["x"]]}
        with pytest.raises(ValueError, match="caption"):
            handler.validate(block, 1, BASE_DIR)

    def test_missing_headers_raises(self, handler):
        block = {"type": "table", "caption": "c", "rows": [["x"]]}
        with pytest.raises(ValueError, match="headers"):
            handler.validate(block, 1, BASE_DIR)

    def test_empty_headers_raises(self, handler):
        block = {"type": "table", "caption": "c", "headers": [], "rows": [["x"]]}
        with pytest.raises(ValueError, match="headers"):
            handler.validate(block, 1, BASE_DIR)

    def test_non_list_headers_raises(self, handler):
        block = {"type": "table", "caption": "c", "headers": "A", "rows": [["x"]]}
        with pytest.raises(ValueError, match="headers"):
            handler.validate(block, 1, BASE_DIR)

    def test_empty_header_string_raises(self, handler):
        block = {"type": "table", "caption": "c", "headers": ["  "], "rows": [["x"]]}
        with pytest.raises(ValueError, match=r"headers\[1\]"):
            handler.validate(block, 1, BASE_DIR)

    def test_missing_rows_raises(self, handler):
        block = {"type": "table", "caption": "c", "headers": ["A"]}
        with pytest.raises(ValueError, match="rows"):
            handler.validate(block, 1, BASE_DIR)

    def test_empty_rows_raises(self, handler):
        block = {"type": "table", "caption": "c", "headers": ["A"], "rows": []}
        with pytest.raises(ValueError, match="rows"):
            handler.validate(block, 1, BASE_DIR)

    def test_row_wrong_column_count_raises(self, handler):
        block = {
            "type": "table",
            "caption": "c",
            "headers": ["A", "B"],
            "rows": [["only one cell"]],
        }
        with pytest.raises(ValueError, match=r"rows\[1\]"):
            handler.validate(block, 1, BASE_DIR)

    def test_non_list_row_raises(self, handler):
        block = {"type": "table", "caption": "c", "headers": ["A"], "rows": ["not a list"]}
        with pytest.raises(ValueError, match=r"rows\[1\]"):
            handler.validate(block, 1, BASE_DIR)

    def test_cell_values_coerced_to_string(self, handler):
        block = {"type": "table", "caption": "c", "headers": ["A"], "rows": [[42]]}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["rows"][0][0] == "42"


class TestTableRender:
    def test_increments_table_counter(self, handler, doc):
        block = {"type": "table", "caption": "cap", "headers": ["A"], "rows": [["x"]]}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        assert ctx.table_counter == 1

    def test_multiple_tables_increment_counter(self, handler, doc):
        block = {"type": "table", "caption": "cap", "headers": ["A"], "rows": [["x"]]}
        ctx = RenderContext()
        for _ in range(3):
            handler.render(doc, block, ctx)
        assert ctx.table_counter == 3

    def test_does_not_affect_figure_counter(self, handler, doc):
        block = {"type": "table", "caption": "cap", "headers": ["A"], "rows": [["x"]]}
        ctx = RenderContext(figure_counter=5)
        handler.render(doc, block, ctx)
        assert ctx.figure_counter == 5

    def test_caption_text_format(self, handler, doc):
        block = {"type": "table", "caption": "Параметры", "headers": ["A"], "rows": [["x"]]}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        # Caption is first — find the GOST Table Caption paragraph
        caption_paras = [p for p in doc.paragraphs if p.style.name == "GOST Table Caption"]
        assert len(caption_paras) == 1
        assert caption_paras[0].text == "Таблица 1 – Параметры"

    def test_caption_strips_trailing_period(self, handler, doc):
        block = {"type": "table", "caption": "Параметры.", "headers": ["A"], "rows": [["x"]]}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        caption_paras = [p for p in doc.paragraphs if p.style.name == "GOST Table Caption"]
        assert caption_paras[0].text == "Таблица 1 – Параметры"

    def test_table_added_to_document(self, handler, doc):
        block = {"type": "table", "caption": "cap", "headers": ["A", "B"], "rows": [["1", "2"]]}
        ctx = RenderContext()
        initial = len(doc.tables)
        handler.render(doc, block, ctx)
        assert len(doc.tables) == initial + 1

    def test_table_dimensions(self, handler, doc):
        block = {
            "type": "table",
            "caption": "cap",
            "headers": ["A", "B", "C"],
            "rows": [["1", "2", "3"], ["4", "5", "6"]],
        }
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        table = doc.tables[-1]
        assert len(table.rows) == 3      # 1 header + 2 data rows
        assert len(table.columns) == 3

    def test_header_row_is_bold(self, handler, doc):
        block = {"type": "table", "caption": "cap", "headers": ["Col"], "rows": [["val"]]}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        header_cell = doc.tables[-1].rows[0].cells[0]
        run = header_cell.paragraphs[0].runs[0]
        assert run.bold is True

    def test_data_row_is_not_bold(self, handler, doc):
        block = {"type": "table", "caption": "cap", "headers": ["Col"], "rows": [["val"]]}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        data_cell = doc.tables[-1].rows[1].cells[0]
        run = data_cell.paragraphs[0].runs[0]
        assert not run.bold

    def test_header_cell_alignment(self, handler, doc):
        block = {"type": "table", "caption": "cap", "headers": ["Col"], "rows": [["val"]]}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        para = doc.tables[-1].rows[0].cells[0].paragraphs[0]
        assert para.alignment == WD_ALIGN_PARAGRAPH.CENTER

    def test_data_cell_alignment(self, handler, doc):
        block = {"type": "table", "caption": "cap", "headers": ["Col"], "rows": [["val"]]}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        para = doc.tables[-1].rows[1].cells[0].paragraphs[0]
        assert para.alignment == WD_ALIGN_PARAGRAPH.LEFT

    def test_sets_current_page_has_content(self, handler, doc):
        block = {"type": "table", "caption": "cap", "headers": ["A"], "rows": [["x"]]}
        ctx = RenderContext(current_page_has_content=False)
        handler.render(doc, block, ctx)
        assert ctx.current_page_has_content is True
