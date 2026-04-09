"""Tests for the paragraph block handler."""

from pathlib import Path

import pytest
from docx import Document
from docx.shared import Cm

from otchet_compose.generator.blocks._base import RenderContext
from otchet_compose.generator.blocks.paragraph import ParagraphHandler
from otchet_compose.generator.styles import setup_styles

BASE_DIR = Path(".")


@pytest.fixture()
def handler():
    return ParagraphHandler()


@pytest.fixture()
def doc():
    d = Document()
    setup_styles(d)
    return d


class TestParagraphValidate:
    def test_valid_block(self, handler):
        block = {"type": "paragraph", "text": "Some text"}
        result = handler.validate(block, 1, BASE_DIR)
        assert result == {"type": "paragraph", "text": "Some text"}

    def test_text_is_stripped(self, handler):
        block = {"type": "paragraph", "text": "  text  "}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["text"] == "text"

    def test_missing_text_raises(self, handler):
        with pytest.raises(ValueError, match="text"):
            handler.validate({"type": "paragraph"}, 1, BASE_DIR)

    def test_empty_text_raises(self, handler):
        with pytest.raises(ValueError, match="text"):
            handler.validate({"type": "paragraph", "text": "   "}, 1, BASE_DIR)

    def test_non_string_text_raises(self, handler):
        with pytest.raises(ValueError, match="text"):
            handler.validate({"type": "paragraph", "text": 42}, 1, BASE_DIR)


class TestParagraphRender:
    def test_adds_paragraph_with_correct_style(self, handler, doc):
        block = {"type": "paragraph", "text": "Hello"}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        last = doc.paragraphs[-1]
        assert last.text == "Hello"
        assert last.style.name == "GOST Body Text"

    def test_first_line_indent(self, handler, doc):
        block = {"type": "paragraph", "text": "Hello"}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        indent = doc.paragraphs[-1].paragraph_format.first_line_indent
        # Cm(1.25) = 450000 EMU; python-docx round-trips through twips so
        # the stored value may differ by up to one twip (~635 EMU).
        assert abs(indent - Cm(1.25)) <= 635

    def test_sets_current_page_has_content(self, handler, doc):
        block = {"type": "paragraph", "text": "Hello"}
        ctx = RenderContext(current_page_has_content=False)
        handler.render(doc, block, ctx)
        assert ctx.current_page_has_content is True
