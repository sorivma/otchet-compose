"""Tests for the heading block handler."""

from pathlib import Path

import pytest
from docx import Document

from otchet_compose.generator.blocks._base import RenderContext
from otchet_compose.generator.blocks.heading import HeadingHandler
from otchet_compose.generator.styles import setup_styles

BASE_DIR = Path(".")


@pytest.fixture()
def handler():
    return HeadingHandler()


@pytest.fixture()
def doc():
    d = Document()
    setup_styles(d)
    return d


class TestHeadingValidate:
    def test_valid_structural(self, handler):
        block = {"type": "heading", "text": "Цель работы", "level": 1, "structural": True}
        result = handler.validate(block, 1, BASE_DIR)
        assert result == {"type": "heading", "text": "Цель работы", "level": 1, "structural": True}

    def test_valid_non_structural(self, handler):
        block = {"type": "heading", "text": "Подраздел", "level": 2, "structural": False}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["structural"] is False

    def test_text_is_stripped(self, handler):
        block = {"type": "heading", "text": "  Заголовок  ", "level": 1, "structural": False}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["text"] == "Заголовок"

    def test_missing_text_raises(self, handler):
        block = {"type": "heading", "level": 1, "structural": True}
        with pytest.raises(ValueError, match="text"):
            handler.validate(block, 1, BASE_DIR)

    def test_empty_text_raises(self, handler):
        block = {"type": "heading", "text": "   ", "level": 1, "structural": True}
        with pytest.raises(ValueError, match="text"):
            handler.validate(block, 1, BASE_DIR)

    def test_invalid_level_raises(self, handler):
        block = {"type": "heading", "text": "Title", "level": 5, "structural": False}
        with pytest.raises(ValueError, match="level"):
            handler.validate(block, 1, BASE_DIR)

    def test_missing_level_raises(self, handler):
        block = {"type": "heading", "text": "Title", "structural": False}
        with pytest.raises(ValueError, match="level"):
            handler.validate(block, 1, BASE_DIR)

    def test_non_bool_structural_raises(self, handler):
        block = {"type": "heading", "text": "Title", "level": 1, "structural": "yes"}
        with pytest.raises(ValueError, match="structural"):
            handler.validate(block, 1, BASE_DIR)

    @pytest.mark.parametrize("level", [1, 2, 3])
    def test_all_valid_levels(self, handler, level):
        block = {"type": "heading", "text": "T", "level": level, "structural": False}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["level"] == level


class TestHeadingRender:
    def test_structural_uppercases_text(self, handler, doc):
        block = {"type": "heading", "text": "цель работы", "level": 1, "structural": True}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        last_para = doc.paragraphs[-1]
        assert last_para.text == "ЦЕЛЬ РАБОТЫ"

    def test_structural_no_page_break_on_first_content(self, handler, doc):
        block = {"type": "heading", "text": "Цель", "level": 1, "structural": True}
        ctx = RenderContext(current_page_has_content=False)
        initial_count = len(doc.paragraphs)
        handler.render(doc, block, ctx)
        # No page break paragraph inserted before
        assert len(doc.paragraphs) == initial_count + 1

    def test_structural_page_break_when_page_has_content(self, handler, doc):
        block = {"type": "heading", "text": "Заключение", "level": 1, "structural": True}
        ctx = RenderContext(current_page_has_content=True)
        initial_count = len(doc.paragraphs)
        handler.render(doc, block, ctx)
        # A page break paragraph + the heading paragraph
        assert len(doc.paragraphs) == initial_count + 2

    def test_non_structural_preserves_case(self, handler, doc):
        block = {"type": "heading", "text": "Постановка задачи", "level": 2, "structural": False}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        assert doc.paragraphs[-1].text == "Постановка задачи"

    def test_sets_current_page_has_content(self, handler, doc):
        block = {"type": "heading", "text": "T", "level": 1, "structural": False}
        ctx = RenderContext(current_page_has_content=False)
        handler.render(doc, block, ctx)
        assert ctx.current_page_has_content is True

    @pytest.mark.parametrize("level,expected_style", [
        (1, "GOST Heading 1"),
        (2, "GOST Heading 2"),
        (3, "GOST Heading 3"),
    ])
    def test_non_structural_style_by_level(self, handler, doc, level, expected_style):
        block = {"type": "heading", "text": "T", "level": level, "structural": False}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        assert doc.paragraphs[-1].style.name == expected_style
