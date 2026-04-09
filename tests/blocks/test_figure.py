"""Tests for the figure block handler."""

from pathlib import Path

import pytest
from docx import Document

from otchet_compose.generator.blocks._base import RenderContext
from otchet_compose.generator.blocks.figure import FigureHandler
from otchet_compose.generator.styles import setup_styles

BASE_DIR = Path(".")


@pytest.fixture()
def handler():
    return FigureHandler()


@pytest.fixture()
def doc():
    d = Document()
    setup_styles(d)
    return d


class TestFigureValidate:
    def test_valid_with_path(self, handler, tmp_path):
        img = tmp_path / "img.png"
        img.write_bytes(b"")
        block = {"type": "figure", "caption": "Some figure", "path": str(img)}
        result = handler.validate(block, 1, tmp_path)
        assert result["caption"] == "Some figure"
        assert result["path"] == str(img.resolve())

    def test_valid_without_path(self, handler):
        block = {"type": "figure", "caption": "Some figure"}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["path"] is None

    def test_caption_is_stripped(self, handler):
        block = {"type": "figure", "caption": "  caption  "}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["caption"] == "caption"

    def test_missing_caption_raises(self, handler):
        with pytest.raises(ValueError, match="caption"):
            handler.validate({"type": "figure"}, 1, BASE_DIR)

    def test_empty_caption_raises(self, handler):
        with pytest.raises(ValueError, match="caption"):
            handler.validate({"type": "figure", "caption": "  "}, 1, BASE_DIR)

    def test_empty_path_string_raises(self, handler):
        with pytest.raises(ValueError, match="path"):
            handler.validate({"type": "figure", "caption": "cap", "path": "  "}, 1, BASE_DIR)

    def test_relative_path_resolved_to_absolute(self, handler, tmp_path):
        block = {"type": "figure", "caption": "cap", "path": "./img.png"}
        result = handler.validate(block, 1, tmp_path)
        assert Path(result["path"]).is_absolute()


class TestFigureRender:
    def test_increments_figure_counter(self, handler, doc):
        block = {"type": "figure", "caption": "cap", "path": None}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        assert ctx.figure_counter == 1

    def test_multiple_figures_increment_counter(self, handler, doc):
        ctx = RenderContext()
        for _ in range(3):
            handler.render(doc, {"type": "figure", "caption": "cap", "path": None}, ctx)
        assert ctx.figure_counter == 3

    def test_caption_format(self, handler, doc):
        block = {"type": "figure", "caption": "Схема стенда", "path": None}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        caption_para = doc.paragraphs[-1]
        assert caption_para.text == "Рисунок 1 – Схема стенда"

    def test_caption_strips_trailing_period(self, handler, doc):
        block = {"type": "figure", "caption": "Схема.", "path": None}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        assert doc.paragraphs[-1].text == "Рисунок 1 – Схема"

    def test_placeholder_inserted_when_no_path(self, handler, doc):
        block = {"type": "figure", "caption": "cap", "path": None}
        ctx = RenderContext()
        initial_tables = len(doc.tables)
        handler.render(doc, block, ctx)
        assert len(doc.tables) == initial_tables + 1

    def test_placeholder_inserted_when_path_missing(self, handler, doc, tmp_path):
        block = {"type": "figure", "caption": "cap", "path": str(tmp_path / "nonexistent.png")}
        ctx = RenderContext()
        initial_tables = len(doc.tables)
        handler.render(doc, block, ctx)
        assert len(doc.tables) == initial_tables + 1

    def test_sets_current_page_has_content(self, handler, doc):
        block = {"type": "figure", "caption": "cap", "path": None}
        ctx = RenderContext(current_page_has_content=False)
        handler.render(doc, block, ctx)
        assert ctx.current_page_has_content is True
