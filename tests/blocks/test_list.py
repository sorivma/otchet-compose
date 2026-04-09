"""Tests for the list block handler."""

from pathlib import Path

import pytest
from docx import Document

from otchet_compose.generator.blocks._base import RenderContext
from otchet_compose.generator.blocks.list import ListHandler
from otchet_compose.generator.styles import setup_styles

BASE_DIR = Path(".")


@pytest.fixture()
def handler():
    return ListHandler()


@pytest.fixture()
def doc():
    d = Document()
    setup_styles(d)
    return d


class TestListValidate:
    def test_valid_bullet(self, handler):
        block = {"type": "list", "style": "bullet", "items": ["A", "B"]}
        result = handler.validate(block, 1, BASE_DIR)
        assert result == {"type": "list", "style": "bullet", "items": ["A", "B"]}

    def test_valid_numeric(self, handler):
        block = {"type": "list", "style": "numeric", "items": ["Step 1", "Step 2"]}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["style"] == "numeric"

    def test_items_are_stripped(self, handler):
        block = {"type": "list", "style": "bullet", "items": ["  item  "]}
        result = handler.validate(block, 1, BASE_DIR)
        assert result["items"] == ["item"]

    def test_invalid_style_raises(self, handler):
        block = {"type": "list", "style": "dashes", "items": ["A"]}
        with pytest.raises(ValueError, match="style"):
            handler.validate(block, 1, BASE_DIR)

    def test_missing_style_raises(self, handler):
        block = {"type": "list", "items": ["A"]}
        with pytest.raises(ValueError, match="style"):
            handler.validate(block, 1, BASE_DIR)

    def test_missing_items_raises(self, handler):
        block = {"type": "list", "style": "bullet"}
        with pytest.raises(ValueError, match="items"):
            handler.validate(block, 1, BASE_DIR)

    def test_empty_items_raises(self, handler):
        block = {"type": "list", "style": "bullet", "items": []}
        with pytest.raises(ValueError, match="items"):
            handler.validate(block, 1, BASE_DIR)

    def test_non_string_item_raises(self, handler):
        block = {"type": "list", "style": "bullet", "items": ["ok", 42]}
        with pytest.raises(ValueError, match=r"items\[2\]"):
            handler.validate(block, 1, BASE_DIR)

    def test_empty_string_item_raises(self, handler):
        block = {"type": "list", "style": "bullet", "items": ["ok", "  "]}
        with pytest.raises(ValueError, match=r"items\[2\]"):
            handler.validate(block, 1, BASE_DIR)


class TestListRender:
    def test_renders_correct_number_of_paragraphs(self, handler, doc):
        block = {"type": "list", "style": "bullet", "items": ["A", "B", "C"]}
        ctx = RenderContext()
        initial = len(doc.paragraphs)
        handler.render(doc, block, ctx)
        assert len(doc.paragraphs) == initial + 3

    def test_paragraph_text_matches_items(self, handler, doc):
        block = {"type": "list", "style": "bullet", "items": ["First", "Second"]}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        texts = [p.text for p in doc.paragraphs[-2:]]
        assert texts == ["First", "Second"]

    def test_paragraph_style(self, handler, doc):
        block = {"type": "list", "style": "numeric", "items": ["Item"]}
        ctx = RenderContext()
        handler.render(doc, block, ctx)
        assert doc.paragraphs[-1].style.name == "GOST List"

    def test_sets_current_page_has_content(self, handler, doc):
        block = {"type": "list", "style": "bullet", "items": ["X"]}
        ctx = RenderContext(current_page_has_content=False)
        handler.render(doc, block, ctx)
        assert ctx.current_page_has_content is True
