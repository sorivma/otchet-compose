"""Top-level document orchestrator.

Wires together page setup, styles, and block rendering into a single
:func:`generate_document` call.  Block-type dispatch is fully delegated
to the :data:`~otchet_compose.generator.blocks.REGISTRY`; this module
contains no block-specific logic.
"""

from pathlib import Path

from docx import Document
from docx.shared import Cm

from .blocks import REGISTRY, RenderContext
from .fields import OxmlHelper
from .page import PageSetup
from .styles import setup_styles
from .title_page import render_title_page


def _remove_initial_empty_paragraph(doc) -> None:
    """Remove the empty paragraph that python-docx inserts into every new Document."""
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text:
        paragraph = doc.paragraphs[0]._element
        paragraph.getparent().remove(paragraph)


def _add_toc(doc) -> None:
    """Insert a Word TOC field preceded by a "СОДЕРЖАНИЕ" title paragraph.

    The field uses ``\\o "1-3" \\h \\z \\u`` switches (levels 1–3, hyperlinks,
    hide tab leaders, use applied paragraph outline levels).  The placeholder
    text reminds users to update fields in Word before printing.
    """
    doc.add_paragraph("СОДЕРЖАНИЕ", style="GOST TOC Title")

    paragraph = doc.add_paragraph(style="GOST Service")
    paragraph.paragraph_format.first_line_indent = Cm(0)

    OxmlHelper.add_field_run(
        paragraph,
        r' TOC \o "1-3" \h \z \u ',
        "Оглавление будет сформировано после ручного обновления полей в Microsoft Word.",
        dirty=False,
    )


def generate_document(config: dict) -> None:
    """Generate a DOCX report from a validated config dict.

    The *config* dict is the value returned by
    :func:`~otchet_compose.config.load_config`.  The output file is
    written to the path given in ``config["document"]["output"]``;
    parent directories are created automatically.
    """
    document_cfg = config["document"]
    content = config["content"]

    doc = Document()

    PageSetup.apply(doc, hide_first_page_number=document_cfg.get("reserve_title_page", False))
    setup_styles(doc)
    _remove_initial_empty_paragraph(doc)

    ctx = RenderContext()

    title_page_cfg = document_cfg.get("title_page")
    if title_page_cfg:
        render_title_page(doc, title_page_cfg["template"], title_page_cfg["params"])
    elif document_cfg.get("reserve_title_page", False):
        doc.add_paragraph("")
        doc.add_page_break()

    if document_cfg["toc"]:
        _add_toc(doc)
        ctx.current_page_has_content = True

    for block in content:
        REGISTRY[block["type"]].render(doc, block, ctx)

    PageSetup.add_page_number(doc.sections[0])

    output_path = Path(document_cfg["output"])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    print(f"Отчёт сохранён: {output_path}")
