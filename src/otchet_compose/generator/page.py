"""Page layout and footer configuration.

Handles A4 dimensions, GOST-style margins, and the page-number field
inserted into the first section's footer.
"""

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt

from .fields import add_field_run, set_font


def setup_page(doc, hide_first_page_number: bool) -> None:
    """Configure the default section for A4 with GOST-style margins.

    Margins: top/bottom 2 cm, left 3 cm, right 1.5 cm.

    Args:
        doc: The python-docx ``Document`` to configure.
        hide_first_page_number: When ``True``, enables a separate
            first-page header/footer so the title page has no page number.
    """
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)
    section.different_first_page_header_footer = hide_first_page_number


def add_page_number(section) -> None:
    """Insert a centred ``PAGE`` field into the footer of *section*."""
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    pf = paragraph.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    run = add_field_run(paragraph, " PAGE ", "1", dirty=False)
    set_font(
        run,
        name="Times New Roman",
        size=14,
        bold=False,
        italic=False,
        underline=False,
        color=(0, 0, 0),
    )