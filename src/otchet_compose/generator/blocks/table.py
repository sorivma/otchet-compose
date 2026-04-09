"""Table block handler.

Validates ``type: table`` blocks and renders them as GOST-styled Word tables.
Each rendered table increments
:attr:`~otchet_compose.generator.blocks._base.RenderContext.table_counter`
so captions are numbered consecutively across the document.

GOST layout:
    - Caption ("Таблица N – text") appears **above** the table, left-aligned.
    - Header row uses bold, centred text.
    - Data cells use normal-weight, left-aligned text.
    - All cell borders: thin (0.5 pt) single-line black.
    - Font: Times New Roman 14 pt throughout.

YAML format::

    - type: table
      caption: "Some caption"
      headers:
        - "Column A"
        - "Column B"
      rows:
        - ["cell 1", "cell 2"]
        - ["cell 3", "cell 4"]
"""

from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ._base import RenderContext

# Single-line border spec: sz=4 ≈ 0.5 pt (half-points)
_BORDER_SZ = "4"


class TableHandler:
    """Handler for ``type: table`` blocks."""

    def validate(self, block: dict, index: int, base_dir: Path) -> dict:
        """Validate a table block and return its normalised form.

        Args:
            block: Raw dict with ``caption``, ``headers``, and ``rows`` keys.
            index: 1-based position in the content list (for error messages).
            base_dir: Config file directory (unused for tables).

        Returns:
            Normalised block dict with ``type``, ``caption``, ``headers``,
            and ``rows``.

        Raises:
            ValueError: If any required field is missing, empty, or has an
                invalid type/shape.
        """
        caption = block.get("caption")
        headers = block.get("headers")
        rows = block.get("rows")

        if not isinstance(caption, str) or not caption.strip():
            raise ValueError(
                f"content[{index}]: table.caption обязателен и должен быть строкой"
            )

        if not isinstance(headers, list) or not headers:
            raise ValueError(
                f"content[{index}]: table.headers обязателен и должен быть непустым списком"
            )

        for i, h in enumerate(headers, start=1):
            if not isinstance(h, str) or not h.strip():
                raise ValueError(
                    f"content[{index}]: table.headers[{i}] должен быть непустой строкой"
                )

        col_count = len(headers)

        if not isinstance(rows, list) or not rows:
            raise ValueError(
                f"content[{index}]: table.rows обязателен и должен быть непустым списком"
            )

        normalised_rows = []
        for r_idx, row in enumerate(rows, start=1):
            if not isinstance(row, list):
                raise ValueError(
                    f"content[{index}]: table.rows[{r_idx}] должен быть списком"
                )
            if len(row) != col_count:
                raise ValueError(
                    f"content[{index}]: table.rows[{r_idx}] содержит {len(row)} ячеек, "
                    f"ожидается {col_count} (по числу заголовков)"
                )
            normalised_rows.append([str(cell) for cell in row])

        return {
            "type": "table",
            "caption": caption.strip(),
            "headers": [h.strip() for h in headers],
            "rows": normalised_rows,
        }

    def render(self, doc, block: dict, ctx: RenderContext) -> None:
        """Append a numbered caption and a bordered GOST table to *doc*."""
        ctx.table_counter += 1

        _add_caption(doc, ctx.table_counter, block["caption"])

        headers = block["headers"]
        rows = block["rows"]
        col_count = len(headers)

        table = doc.add_table(rows=1 + len(rows), cols=col_count)
        table.autofit = True
        _set_table_borders(table)

        # Header row
        header_row = table.rows[0]
        for col_idx, text in enumerate(headers):
            cell = header_row.cells[col_idx]
            para = cell.paragraphs[0]
            para.style = "GOST Table Cell"
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(text)
            run.bold = True

        # Data rows
        for row_idx, row_data in enumerate(rows, start=1):
            word_row = table.rows[row_idx]
            for col_idx, text in enumerate(row_data):
                cell = word_row.cells[col_idx]
                para = cell.paragraphs[0]
                para.style = "GOST Table Cell"
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.add_run(text)

        ctx.current_page_has_content = True


# ---------------------------------------------------------------------------
# Private helpers
# ---------------------------------------------------------------------------

def _add_caption(doc, table_number: int, text: str) -> None:
    """Append a ``GOST Table Caption`` paragraph formatted as "Таблица N – text"."""
    caption_text = f"Таблица {table_number} – {text.strip().rstrip('.')}"
    doc.add_paragraph(caption_text, style="GOST Table Caption")


def _set_table_borders(table) -> None:
    """Apply a thin black single-line border to all sides and inner grid of *table*."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{side}"))
        if border is None:
            border = OxmlElement(f"w:{side}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), _BORDER_SZ)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
