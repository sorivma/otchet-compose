"""Figure block handler.

Validates ``type: figure`` blocks and renders them as either an embedded
image (when ``path`` points to an existing file) or a bordered placeholder
inside a two-row table together with the caption so the block does not
split across pages.
"""

from pathlib import Path

from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from ._base import RenderContext


class FigureHandler:
    """Handler for ``type: figure`` blocks."""

    def validate(self, block: dict, index: int, base_dir: Path) -> dict:
        caption = block.get("caption")
        path = block.get("path")

        if not isinstance(caption, str) or not caption.strip():
            raise ValueError(
                f"content[{index}]: figure.caption обязателен и должен быть строкой"
            )

        normalized_path = None
        if path is not None:
            if not isinstance(path, str) or not path.strip():
                raise ValueError(
                    f"content[{index}]: figure.path должен быть непустой строкой, если указан"
                )
            normalized_path = str(_resolve_path(base_dir, path))

        return {
            "type": "figure",
            "caption": caption.strip(),
            "path": normalized_path,
        }

    def render(self, doc, block: dict, ctx: RenderContext) -> None:
        """Append a figure and caption as one non-breaking block."""
        ctx.figure_counter += 1
        image_path = block.get("path")
        caption_text = _caption_text(ctx.figure_counter, block["caption"])

        table = doc.add_table(rows=2, cols=1)
        table.autofit = False
        _clear_table_borders(table)

        image_cell = table.cell(0, 0)
        caption_cell = table.cell(1, 0)
        image_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        caption_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        image_row = table.rows[0]
        if image_path and Path(image_path).exists():
            image_para = image_cell.paragraphs[0]
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_para.add_run().add_picture(image_path, width=Cm(16))
        else:
            image_row.height = Cm(7)
            image_row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            _set_cell_borders(image_cell)
            image_para = image_cell.paragraphs[0]
            image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            image_para.style = "GOST Figure Placeholder"
            image_para.add_run("Изображение отсутствует")

        caption_para = caption_cell.paragraphs[0]
        caption_para.style = "GOST Caption"
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_para.add_run(caption_text)

        _set_row_cant_split(image_row)
        _set_row_cant_split(table.rows[1])
        ctx.current_page_has_content = True


def _caption_text(figure_number: int, text: str) -> str:
    return f"Рисунок {figure_number} – {text.strip().rstrip('.')}"


def _set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def _clear_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")


def _set_cell_borders(cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def _resolve_path(base_dir: Path, raw_path: str) -> Path:
    path = Path(raw_path)
    if path.is_absolute():
        return path.resolve()
    return (base_dir / path).resolve()
