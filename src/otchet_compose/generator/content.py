from pathlib import Path

from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

from .fields import add_field_run, set_table_borders


def add_toc(doc) -> None:
    doc.add_paragraph("СОДЕРЖАНИЕ", style="GOST TOC Title")

    paragraph = doc.add_paragraph(style="GOST Service")
    paragraph.paragraph_format.first_line_indent = Cm(0)

    add_field_run(
        paragraph,
        r' TOC \o "1-3" \h \z \u ',
        "Оглавление будет сформировано после ручного обновления полей в Microsoft Word.",
        dirty=False,
    )


def add_heading(doc, text: str, level: int, structural: bool, page_break_before: bool = False) -> None:
    text = text.strip()

    if structural:
        if page_break_before:
            doc.add_page_break()

        style_name = "GOST Heading 1" if level == 1 else "GOST Heading 2"
        doc.add_paragraph(text.upper(), style=style_name)
        return

    style_name = {
        1: "GOST Heading 1",
        2: "GOST Heading 2",
        3: "GOST Heading 3",
    }[level]
    doc.add_paragraph(text, style=style_name)


def add_body_paragraph(doc, text: str) -> None:
    paragraph = doc.add_paragraph(text.strip(), style="GOST Body Text")
    paragraph.paragraph_format.first_line_indent = Cm(1.25)


def add_figure(doc, caption: str, image_path: str | None, figure_number: int) -> None:
    if image_path and Path(image_path).exists():
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        run.add_picture(image_path, width=Cm(16))
    else:
        add_figure_placeholder(doc)

    add_figure_caption(doc, figure_number, caption)


def add_figure_placeholder(doc) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    set_table_borders(table)

    cell = table.cell(0, 0)
    cell.width = Cm(12)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    row = table.rows[0]
    row.height = Cm(7)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.style = "GOST Figure Placeholder"
    paragraph.add_run("Изображение отсутствует")


def add_figure_caption(doc, figure_number: int, text: str) -> None:
    caption_text = f"Рисунок {figure_number} – {text.strip().rstrip('.')}"
    doc.add_paragraph(caption_text, style="GOST Caption")