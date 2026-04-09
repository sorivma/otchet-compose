"""Built-in title page template: МИРЭА – Российский технологический университет.

Supported params (all optional — omit to leave blank):

    institute   Institute name, e.g. "Институт информационных технологий"
    department  Department name, e.g. "Кафедра общей информатики"
    discipline  Subject name, e.g. "Операционные системы"
    work_type   Work type, e.g. "ОТЧЁТ ПО ЛАБОРАТОРНОЙ РАБОТЕ"
    work_number Work number, e.g. "№ 4"
    work_title  Optional title of the work
    student     Full name of student, e.g. "Петров П.П."
    group       Group code, e.g. "ИКБО-01-22"
    teacher     Full name of teacher, e.g. "Иванов И.И."
    year        Year, e.g. "2025" (defaults to current year if omitted)
"""

from datetime import date

from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Cm, Pt

from otchet_compose.generator.fields import OxmlHelper


def render(doc, params: dict) -> None:
    """Render the MIREA title page into *doc*."""
    year = params.get("year", str(date.today().year))

    _p(doc, "МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ",
       align="center", size=12, space_before=0)
    _p(doc,
       "Федеральное государственное бюджетное образовательное учреждение "
       "высшего образования",
       align="center", size=12)
    _p(doc, "«МИРЭА – Российский технологический университет»",
       align="center", size=12, bold=True)
    _p(doc, "РТУ МИРЭА", align="center", size=14, bold=True, space_before=6)

    institute = params.get("institute", "")
    if institute:
        _p(doc, institute, align="center", size=12, space_before=6)

    department = params.get("department", "")
    if department:
        _p(doc, department, align="center", size=12)

    _separator(doc)

    work_type = params.get("work_type", "ОТЧЁТ ПО ЛАБОРАТОРНОЙ РАБОТЕ")
    _p(doc, work_type, align="center", size=14, bold=True, space_before=0)

    work_number = params.get("work_number", "")
    if work_number:
        _p(doc, work_number, align="center", size=14, bold=True)

    discipline = params.get("discipline", "")
    if discipline:
        _p(doc, f"по дисциплине «{discipline}»", align="center", size=12, space_before=6)

    work_title = params.get("work_title", "")
    if work_title:
        _p(doc, f"«{work_title}»", align="center", size=12, bold=True)

    _separator(doc)

    student = params.get("student", "")
    group = params.get("group", "")
    teacher = params.get("teacher", "")

    _p(doc, "", space_before=0)  # spacer

    if student or group:
        student_line = "Выполнил:"
        if group:
            student_line += f"  студент группы {group}"
        _p_right_block(doc, student_line, params.get("student", ""))

    if teacher:
        _p_right_block(doc, "Принял:", teacher)

    # Push "Москва YYYY" to the bottom with spacing
    for _ in range(4):
        _p(doc, "")

    _p(doc, f"Москва  {year}", align="center", size=12, space_before=0)

    doc.add_page_break()


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _p(doc, text: str, *, align: str = "left", size: int = 14,
       bold: bool = False, space_before: int = 0) -> None:
    """Add a single paragraph with Times New Roman formatting."""
    para = doc.add_paragraph()
    para.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]

    pf = para.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    if text:
        run = para.add_run(text)
        OxmlHelper.set_font(run, size=size, bold=bold)


def _p_right_block(doc, label: str, value: str) -> None:
    """Add a right-aligned label/value pair (e.g. 'Выполнил:  Петров П.П.')."""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    pf = para.paragraph_format
    pf.first_line_indent = Cm(0)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.space_before = Pt(6)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE

    run = para.add_run(f"{label}  {value}" if value else label)
    OxmlHelper.set_font(run, size=12)


def _separator(doc) -> None:
    """Add a blank paragraph acting as a visual separator."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.first_line_indent = Cm(0)
