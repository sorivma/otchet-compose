"""Title page renderer.

Templates are ``.docx`` files with ``{{key}}`` placeholders in their text.
Two directories are searched, in order:

1. ``~/.otchet-compose/templates/<name>.docx``  — user-defined (checked first)
2. ``<package>/generator/templates/<name>.docx`` — bundled with the tool

User templates always shadow bundled ones of the same name.  To inspect or
customise a bundled template, copy it from the package directory to
``~/.otchet-compose/templates/`` and edit it in Word.

Placeholder substitution
------------------------
``{{key}}`` tokens are replaced by the matching value from *params*.
Missing keys are left unchanged.  Tokens that Word has split across multiple
``Run`` objects are handled by collapsing the paragraph text into the first
run before substitution.

Merging
-------
All body elements of the template (except its ``w:sectPr``) are prepended
to the output document.  Styles and relationships (e.g. images) from the
template are **not** transferred — template authors should use only standard
Word styles and avoid embedded images.
"""

import copy
import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .templates import BUNDLED_TEMPLATES_DIR

_USER_TEMPLATES_DIR = Path.home() / ".otchet-compose" / "templates"
_PLACEHOLDER_RE = re.compile(r"\{\{([^}]+)\}\}")


def render_title_page(doc, name: str, params: dict) -> None:
    """Render the named title page template as the first content of *doc*.

    Args:
        doc: The python-docx ``Document`` being built.
        name: Template name without the ``.docx`` extension.
        params: Mapping of placeholder keys to replacement strings.

    Raises:
        ValueError: If no template with *name* is found in either location.
    """
    path = _locate_template(name)
    tmpl = Document(str(path))
    _substitute_all(tmpl, params)
    _prepend_body(doc, tmpl)


def list_available_templates() -> list[str]:
    """Return sorted names of all available templates (user + bundled, deduplicated)."""
    names: set[str] = {p.stem for p in BUNDLED_TEMPLATES_DIR.glob("*.docx")}
    if _USER_TEMPLATES_DIR.exists():
        names.update(p.stem for p in _USER_TEMPLATES_DIR.glob("*.docx"))
    return sorted(names)


# ---------------------------------------------------------------------------
# Template lookup
# ---------------------------------------------------------------------------

def _locate_template(name: str) -> Path:
    """Return the path to the *name* template, preferring user over bundled.

    Raises:
        ValueError: If no matching ``.docx`` file is found.
    """
    user_path = _USER_TEMPLATES_DIR / f"{name}.docx"
    if user_path.exists():
        return user_path

    bundled_path = BUNDLED_TEMPLATES_DIR / f"{name}.docx"
    if bundled_path.exists():
        return bundled_path

    available = list_available_templates()
    hint = ", ".join(available) if available else "нет"
    raise ValueError(
        f"Шаблон титульного листа {name!r} не найден. "
        f"Доступные шаблоны: {hint}. "
        f"Чтобы добавить собственный шаблон, поместите "
        f"{name}.docx в {_USER_TEMPLATES_DIR}."
    )


# ---------------------------------------------------------------------------
# Placeholder substitution
# ---------------------------------------------------------------------------

def _substitute_all(tmpl_doc, params: dict) -> None:
    """Replace ``{{key}}`` in every paragraph of *tmpl_doc*, including tables."""
    for para in _iter_all_paragraphs(tmpl_doc):
        _substitute_paragraph(para, params)


def _iter_all_paragraphs(doc):
    """Yield every paragraph in *doc*, including paragraphs inside table cells."""
    yield from doc.paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs


def _substitute_paragraph(para, params: dict) -> None:
    """Substitute ``{{key}}`` in *para*, handling tokens split across runs.

    Word may internally split a token like ``{{name}}`` across several ``Run``
    objects.  This method collapses all run texts into the first run (keeping
    its character formatting), performs substitution, then clears the rest.
    """
    full_text = "".join(run.text for run in para.runs)
    if "{{" not in full_text:
        return

    substituted = _PLACEHOLDER_RE.sub(
        lambda m: params.get(m.group(1).strip(), m.group(0)),
        full_text,
    )

    if para.runs:
        para.runs[0].text = substituted
        for run in para.runs[1:]:
            run.text = ""


# ---------------------------------------------------------------------------
# Document merging
# ---------------------------------------------------------------------------

def _prepend_body(doc, tmpl_doc) -> None:
    """Insert all non-sectPr body elements of *tmpl_doc* at the start of *doc*.

    A page break paragraph is appended after the template content so that the
    report body always starts on a fresh page regardless of how the template
    file ends.  This relieves template authors from having to remember to add
    a manual page break at the bottom of their file.
    """
    body = doc.element.body
    children = list(body)
    insert_before = children[0] if children else None

    def _insert(elem):
        if insert_before is not None:
            body.insert(list(body).index(insert_before), elem)
        else:
            body.append(elem)

    for elem in tmpl_doc.element.body:
        if elem.tag == qn("w:sectPr"):
            continue
        _insert(copy.deepcopy(elem))

    # Guarantee a page break between the title page and the report body.
    _insert(_make_page_break_paragraph())


def _make_page_break_paragraph():
    """Return a ``w:p`` element containing a ``w:lastRenderedPageBreak``-style
    page break run (``w:br`` with ``w:type="page"``)."""
    p_elem = OxmlElement("w:p")
    r_elem = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r_elem.append(br)
    p_elem.append(r_elem)
    return p_elem
